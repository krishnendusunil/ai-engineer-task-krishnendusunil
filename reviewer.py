import os
import glob
import pickle
import faiss
from pathlib import Path
from docx import Document
from docx.shared import RGBColor
from sentence_transformers import SentenceTransformer
import fitz  # PyMuPDF for PDFs
from dotenv import load_dotenv
import json
import google.generativeai as genai
from tabulate import tabulate
from colorama import Fore, Style, init

# --- Load environment variables ---
load_dotenv()

# --- Settings ---
FAISS_PATH = "faiss.index"
TEXTS_PATH = "texts.pkl"
MODEL_NAME = "all-MiniLM-L6-v2"
UPLOADS_DIR = "uploads"
REVIEWED_DIR = "reviewed"
REPORTS_DIR = "reports"

# Load FAISS + text chunks
index = faiss.read_index(FAISS_PATH)
data = pickle.load(open(TEXTS_PATH, "rb"))

# Embedding model
embedder = SentenceTransformer(MODEL_NAME)

# Configure Gemini API
google_api_key = os.getenv("GOOGLE_API_KEY")
if not google_api_key:
    raise ValueError("Missing GOOGLE_API_KEY in .env")
genai.configure(api_key=google_api_key)
gemini_model = genai.GenerativeModel("gemini-1.5-flash")

# Initialize colorama
init(autoreset=True)


# ---------- Simple doc type classifier & checklists ----------
# Keyword-based doc type detection (expand as needed)
DOC_TYPE_KEYWORDS = {
    "Articles of Association": ["articles of association", "articles of association (", "articles of assoc"],
    "Memorandum of Association": ["memorandum of association", "memorandum of assoc", "moa"],
    "Board Resolution": ["board resolution", "resolution of the board", "board of directors resolution"],
    "Shareholder Resolution": ["shareholder resolution", "resolution of the shareholders"],
    "Incorporation Application Form": ["application for incorporation", "incorporation application"],
    "UBO Declaration Form": ["ubo declaration", "ultimate beneficial owner", "ubo"],
    "Register of Members and Directors": ["register of members", "register of directors", "register of members and directors"],
    "Change of Registered Address Notice": ["change of registered address", "registered address notice"],
    "Employment Contract": ["standard employment contract", "employment contract"],
    "Data Protection Policy": ["appropriate policy document", "data protection"],
}

# Required docs checklist per process
REQUIRED_DOCS = {
    "Company Incorporation": [
        "Articles of Association",
        "Memorandum of Association",
        "Incorporation Application Form",
        "UBO Declaration Form",
        "Register of Members and Directors"
    ],
    # you can add more processes and required doc lists here
}


def detect_doc_type(text):
    txt = (text or "").lower()
    for dtype, keywords in DOC_TYPE_KEYWORDS.items():
        for kw in keywords:
            if kw in txt:
                return dtype
    return "Unknown"


def detect_process(uploaded_types):
    # Simple rules: if AoA or MoA present => Company Incorporation
    if any(t in ["Articles of Association", "Memorandum of Association", "Incorporation Application Form"] for t in uploaded_types):
        return "Company Incorporation"
    # fallback
    return "Unknown"


# ---------- File helpers ----------
def read_docx(path):
    doc = Document(path)
    return doc, "\n".join([p.text for p in doc.paragraphs])


def read_pdf(path):
    text = ""
    with fitz.open(path) as pdf:
        for page in pdf:
            text += page.get_text()
    return None, text


def retrieve_context(query_text, top_k=5):
    q_emb = embedder.encode([query_text])
    D, I = index.search(q_emb, top_k)
    hits = []
    for idx in I[0]:
        hits.append({
            "text": data["texts"][idx],
            "meta": data["metadatas"][idx]
        })
    return hits


def add_inline_annotation(doc, para_index, comment):
    if doc is None:
        return
    if 0 <= para_index < len(doc.paragraphs):
        run = doc.paragraphs[para_index].add_run(f"  [REVIEW: {comment}]")
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)


# ---------- LLM + parsing ----------
def call_gemini_for_review(context_chunks, doc_text):
    context_text = "\n---\n".join([c["text"] for c in context_chunks])
    prompt = f"""
You are an expert ADGM compliance reviewer.
Use the ADGM reference excerpts below to identify:
- Missing or incorrect clauses
- Wrong jurisdiction references
- Missing required sections/signatures
- Ambiguous language
- Non-compliance with ADGM templates

ADGM REFERENCE MATERIAL:
{context_text}

DOCUMENT TO REVIEW:
{doc_text}

Return a JSON array of issues with fields:
document_section (approx), issue, severity (Low/Medium/High), suggestion, source_reference.
"""
    gemini_response = gemini_model.generate_content(prompt)
    output_text = gemini_response.text.strip()

    # remove ```json ... ``` wrappers if present
    if output_text.startswith("```"):
        output_text = output_text.strip("`")
        if output_text.lower().startswith("json"):
            output_text = output_text[4:].strip()

    try:
        issues = json.loads(output_text)
    except json.JSONDecodeError:
        issues = [{"error": "LLM did not return valid JSON", "raw": output_text}]
    return issues


# ---------- Single-file analysis ----------
def analyze_single_file(path, output_dir=REVIEWED_DIR):
    ext = path.lower()
    if ext.endswith(".docx"):
        doc, doc_text = read_docx(path)
    elif ext.endswith(".pdf"):
        doc, doc_text = read_pdf(path)
    else:
        raise ValueError("Unsupported file type: " + path)

    doc_type = detect_doc_type(doc_text)
    context_chunks = retrieve_context(doc_text, top_k=5)
    issues = call_gemini_for_review(context_chunks, doc_text)

    # Add paragraph index heuristics: if an issue mentions clause text, try to find paragraph
    for issue in issues:
        # attempt to set paragraph_index if not provided
        if "paragraph_index" not in issue:
            # naive: search for the section text in paragraphs
            if doc is not None:
                found_idx = 0
                q = issue.get("document_section", "")
                qlow = (q or "").lower()
                for i, p in enumerate(doc.paragraphs):
                    if qlow and qlow in p.text.lower():
                        found_idx = i
                        break
                issue["paragraph_index"] = found_idx

    # annotate docx
    reviewed_path = None
    if doc is not None:
        Path(output_dir).mkdir(exist_ok=True)
        reviewed_path = Path(output_dir) / f"reviewed_{Path(path).name}"
        for issue in issues:
            idx = issue.get("paragraph_index", 0)
            add_inline_annotation(doc, idx, issue.get("suggestion", issue.get("issue", "")))
        doc.save(reviewed_path)

    return {
        "file": path,
        "doc_type": doc_type,
        "reviewed_file": str(reviewed_path) if reviewed_path else None,
        "issues": issues
    }


# ---------- Batch processing, checklist verification ----------
def analyze_uploads_folder(uploads_dir=UPLOADS_DIR):
    # find .docx in uploads (task requires .docx primarily)
    files = sorted(glob.glob(os.path.join(uploads_dir, "*.docx")))
    if not files:
        raise FileNotFoundError(f"No .docx files found in '{uploads_dir}'")

    results = []
    uploaded_types = []

    for f in files:
        res = analyze_single_file(f)
        results.append(res)
        dtype = res.get("doc_type", "Unknown")
        uploaded_types.append(dtype)

    process = detect_process(uploaded_types)
    required = REQUIRED_DOCS.get(process, [])
    missing = [r for r in required if r not in uploaded_types]

    combined_output = {
        "process": process,
        "documents_uploaded": len(files),
        "required_documents": len(required),
        "missing_documents": missing,
        "per_file_results": results
    }

    # save combined JSON
    Path(REPORTS_DIR).mkdir(exist_ok=True)
    out_json_path = Path(REPORTS_DIR) / "combined_review_report.json"
    with open(out_json_path, "w", encoding="utf-8") as fh:
        json.dump(combined_output, fh, indent=2, ensure_ascii=False)

    return combined_output


# ---------- Console table ----------
def print_results_table_single(file_result):
    table = []
    for issue in file_result["issues"]:
        sev_raw = issue.get("severity", "")
        sev = (sev_raw or "").lower()
        if sev == "high":
            sev_color = Fore.RED + issue.get("severity", "") + Style.RESET_ALL
        elif sev == "medium":
            sev_color = Fore.YELLOW + issue.get("severity", "") + Style.RESET_ALL
        elif sev == "low":
            sev_color = Fore.GREEN + issue.get("severity", "") + Style.RESET_ALL
        else:
            sev_color = issue.get("severity", "")

        table.append([
            issue.get("document_section", ""),
            issue.get("issue", ""),
            sev_color,
            issue.get("suggestion", ""),
            issue.get("source_reference", "")
        ])

    headers = ["Section", "Issue", "Severity", "Suggestion", "Reference"]
    print(f"\nResults for: {file_result['file']} (type: {file_result.get('doc_type')})")
    print(tabulate(table, headers=headers, tablefmt="fancy_grid"))


def print_combined_report(report):
    print("\n====== CHECKLIST SUMMARY ======")
    print(f"Detected process: {report['process']}")
    print(f"Uploaded documents: {report['documents_uploaded']}")
    print(f"Required documents for process: {report['required_documents']}")
    if report["missing_documents"]:
        print(Fore.RED + f"Missing documents: {report['missing_documents']}" + Style.RESET_ALL)
    else:
        print(Fore.GREEN + "No required documents missing." + Style.RESET_ALL)

    # print per-file tables
    for file_res in report["per_file_results"]:
        print_results_table_single(file_res)


# ---------- CLI entry ----------
if __name__ == "__main__":
    combined = analyze_uploads_folder(UPLOADS_DIR)
    print_combined_report(combined)
    print(f"\nCombined JSON saved to: {Path(REPORTS_DIR)/'combined_review_report.json'}")
