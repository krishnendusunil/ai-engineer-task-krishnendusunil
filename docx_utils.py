# docx_utils.py
from docx import Document
from docx.shared import RGBColor

def parse_docx(path):
    doc = Document(path)
    paragraphs = [p.text for p in doc.paragraphs]
    return doc, paragraphs

def add_inline_annotation(doc, para_index, text):
    # simple: append a red bold run at end of paragraph
    from docx.shared import RGBColor
    p = doc.paragraphs[para_index]
    r = p.add_run(f"  [REVIEW: {text}]")
    r.bold = True
    try:
        r.font.color.rgb = RGBColor(0xFF,0x00,0x00)
    except Exception:
        pass

def save_doc(doc, out_path):
    doc.save(out_path)
